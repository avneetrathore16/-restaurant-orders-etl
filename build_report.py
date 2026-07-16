from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(__file__).resolve().parent / "Restaurant_Orders_ETL_Project_Report.docx"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GRAY = "666666"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold=False, italic=False, align=None, after=6, size=11):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    r.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_font(r)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_widths(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(text))
        set_font(r, size=9, bold=True, color=BLUE)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_code(doc, path: Path, label: str):
    heading(doc, label, 2)
    add_text(doc, f"Complete source listing: {path.name}", italic=True, size=9, after=4)
    lines = path.read_text(encoding="utf-8").splitlines()
    for number, line in enumerate(lines, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.left_indent = Inches(.08)
        number_run = p.add_run(f"{number:>3}  ")
        set_font(number_run, size=7.2, color="808080", name="Courier New")
        code_run = p.add_run(line or " ")
        set_font(code_run, size=7.2, name="Courier New")


def setup(doc):
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = Inches(.49)
    section.footer_distance = Inches(.49)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Restaurant Orders Data Curation | Project Report")
    set_font(r, size=8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Restaurant Orders ETL Pipeline and Streamlit Dashboard")
    set_font(r, size=8, color=GRAY)


def page_break(doc):
    doc.add_page_break()


def main():
    doc = Document()
    setup(doc)

    # Cover page
    for _ in range(6):
        doc.add_paragraph()
    add_text(doc, "PROJECT REPORT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, after=18)
    add_text(doc, "Restaurant Orders Data Curation", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=27, after=6)
    add_text(doc, "Using Python ETL Pipeline and Streamlit Dashboard", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, after=36)
    add_text(doc, "Submitted in partial fulfillment of the requirements for the award of", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=6)
    add_text(doc, "[DEGREE / COURSE NAME]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=30)
    add_text(doc, "Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=3)
    add_text(doc, "[STUDENT NAME]  |  [REGISTER NUMBER]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=24)
    add_text(doc, "Under the guidance of", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=3)
    add_text(doc, "[GUIDE NAME]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=34)
    add_text(doc, "[DEPARTMENT NAME]\n[COLLEGE / UNIVERSITY NAME]\nAcademic Year 2025–2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=0)
    page_break(doc)

    heading(doc, "Certificate")
    add_text(doc, "This is to certify that the project report entitled “Restaurant Orders Data Curation Using Python ETL Pipeline and Streamlit Dashboard” is a bona fide work carried out by [STUDENT NAME], [REGISTER NUMBER], in partial fulfillment of the requirements for [DEGREE / COURSE NAME] during the academic year 2025–2026.", after=18)
    add_text(doc, "The work embodied in this report has been completed under my supervision and guidance.", after=48)
    add_text(doc, "Project Guide: ____________________                         Head of Department: ____________________", after=28)
    add_text(doc, "Date: ____________________                                           Place: ____________________")
    page_break(doc)

    heading(doc, "Acknowledgement")
    add_text(doc, "I express my sincere gratitude to my project guide, faculty members, and institution for their guidance and encouragement throughout this project. I also thank my peers and family for their support. This work provided an opportunity to apply data engineering concepts—data cleaning, validation, transformation, integration, storage, and visualization—to a realistic restaurant-order dataset.")
    page_break(doc)

    heading(doc, "Abstract")
    add_text(doc, "This project develops an end-to-end data curation solution for restaurant orders using Python, Pandas, SQLite, and Streamlit. The source consists of an Orders table containing transactional details and a Restaurants table containing restaurant attributes. The solution extracts both files, standardizes columns and text fields, converts dates and numeric values, removes duplicate identifiers, validates business rules, joins the datasets, engineers analytical features, and loads curated data into both CSV and SQLite formats.")
    add_text(doc, "The resulting dashboard provides KPI cards and interactive views for order volume, revenue, payment modes, order amounts, delivery performance, cuisine, and zone. The validation run processed 500 orders and 20 restaurants, producing 500 curated rows with no rejected records, duplicate order IDs, unmatched restaurants, or remaining missing values. The project demonstrates a reproducible, practical workflow that improves raw operational data into trusted analytical data.")
    add_text(doc, "Keywords: ETL, data curation, data validation, Pandas, SQLite, Streamlit, restaurant analytics.", italic=True)
    page_break(doc)

    heading(doc, "Table of Contents")
    contents = ["Certificate", "Acknowledgement", "Abstract", "Chapter 1  Introduction", "Chapter 2  Literature Review", "Chapter 3  System Analysis", "Chapter 4  System Design", "Chapter 5  Dataset Description", "Chapter 6  Methodology", "Chapter 7  Implementation", "Chapter 8  Dashboard", "Chapter 9  Testing", "Chapter 10  Results", "Chapter 11  Conclusion", "Chapter 12  Future Scope", "References", "Appendices"]
    for i, item in enumerate(contents, 1):
        add_text(doc, f"{i:02d}.  {item}", after=3)
    page_break(doc)

    heading(doc, "Chapter 1  Introduction")
    add_text(doc, "Food delivery and restaurant platforms create a continuous stream of operational data. However, raw files are not immediately suitable for decision-making because they may contain inconsistent labels, formatting issues, duplicate records, missing attributes, and weak validation. A structured Extract–Transform–Load (ETL) process converts such data into a dependable analytical asset.")
    heading(doc, "1.1 Problem Statement", 2)
    add_text(doc, "The Orders and Restaurants datasets are stored separately and use different naming conventions for their restaurant key. Without data cleaning and integration, it is difficult to analyze restaurant-level performance, payment behavior, delivery speed, ratings, cuisine patterns, or zones accurately.")
    heading(doc, "1.2 Objectives", 2)
    add_bullets(doc, ["Build a reproducible ETL pipeline for restaurant-order data.", "Clean, standardize, and validate raw source records.", "Integrate transactional orders with restaurant metadata.", "Create meaningful derived features for analysis.", "Load curated data to CSV and SQLite.", "Present results through an interactive Streamlit dashboard."])
    heading(doc, "1.3 Scope", 2)
    add_text(doc, "The scope is limited to the supplied files and their 2022 order records. The system supports batch processing, structured outputs, and exploratory dashboard analysis. It does not perform real-time order ingestion, predictive modeling, or external API integration.")
    page_break(doc)

    heading(doc, "Chapter 2  Literature Review")
    add_text(doc, "ETL is a foundational data engineering pattern in which information is extracted from source systems, transformed to meet data-quality and analytical requirements, and loaded into a target store. Kimball and Caserta describe ETL as a critical step in preparing data for reliable business intelligence. In practical analytics work, careful data cleaning is also essential because real-world datasets commonly contain type inconsistencies, missing values, and duplicate or invalid records.")
    add_text(doc, "Pandas is widely used for tabular data preparation in Python because it supports file ingestion, type conversion, filtering, aggregation, joining, and export. SQLite offers a lightweight relational database for portable storage and SQL-based inspection. Streamlit supports rapid creation of data applications in Python, allowing an analytical dashboard to be built directly on curated data. These tools form a suitable low-cost stack for an academic data curation project.")
    heading(doc, "2.1 Gap Addressed", 2)
    add_text(doc, "Many small datasets are analyzed directly in spreadsheets or notebooks without an explicit validation and load stage. This project addresses that gap by making each stage repeatable: raw inputs are preserved, quality rules are programmed, the merge is explicit, outputs are persisted, and dashboard views use only curated data.")
    page_break(doc)

    heading(doc, "Chapter 3  System Analysis")
    heading(doc, "3.1 Existing System", 2)
    add_text(doc, "In a manual approach, users inspect CSV files separately, clean values in a spreadsheet, merge data manually, and create charts repeatedly. This is time-consuming, difficult to audit, and prone to inconsistent results when source files change.")
    heading(doc, "3.2 Proposed System", 2)
    add_text(doc, "The proposed system is a Python-based batch ETL pipeline. It discovers input files, cleans them consistently, enforces quality rules, creates features, and writes auditable outputs. A Streamlit dashboard consumes the curated CSV and provides interactive filters and visual summaries.")
    add_table(doc, ["Aspect", "Manual approach", "Proposed system"], [["Cleaning", "Repeated and manual", "Programmatic and repeatable"], ["Validation", "Visual inspection", "Defined business rules"], ["Integration", "Spreadsheet merge", "Many-to-one data join"], ["Output", "Ad-hoc sheets", "CSV, SQLite, JSON report"], ["Reporting", "Static charts", "Interactive dashboard"]], [1.25, 2.45, 2.8])
    heading(doc, "3.3 Functional Requirements", 2)
    add_bullets(doc, ["Read Orders.csv and Restaurants.csv.", "Normalize fields, dates, numerical values, and payment labels.", "Validate positive quantities, amounts, and delivery time; ratings from 1 to 5.", "Merge records by restaurant_id.", "Create features and save curated outputs.", "Provide dashboard filters, KPIs, charts, and CSV download."])
    page_break(doc)

    heading(doc, "Chapter 4  System Design")
    heading(doc, "4.1 Architecture", 2)
    add_text(doc, "The design follows a clear batch flow: Raw Orders.csv + Raw Restaurants.csv → Extract → Clean and Validate → Merge → Feature Engineering → Curated CSV / SQLite / Validation JSON → Streamlit Dashboard.")
    heading(doc, "4.2 Data Flow", 2)
    add_table(doc, ["Stage", "Input", "Processing", "Output"], [["Extract", "Two CSV files", "Read with Pandas", "Raw data frames"], ["Transform", "Raw data frames", "Clean, validate, standardize", "Validated records"], ["Integrate", "Orders + restaurants", "Left many-to-one join", "Curated dataset"], ["Load", "Curated dataset", "CSV and SQLite write", "Analytical stores"], ["Present", "Curated CSV", "Filters and visualizations", "Dashboard insights"]], [1.0, 1.45, 2.3, 1.75])
    heading(doc, "4.3 Data Model", 2)
    add_text(doc, "Orders is the transaction-level entity. Each order references a restaurant through restaurant_id. Restaurants is the master entity and contains a unique restaurant_id, restaurant_name, cuisine, zone, and category. The merged dataset retains the transaction grain while adding restaurant attributes.")
    page_break(doc)

    heading(doc, "Chapter 5  Dataset Description")
    add_text(doc, "The project uses two related comma-separated files. The raw order source contains 500 records, and the restaurant source contains 20 records.")
    add_table(doc, ["File", "Rows", "Columns", "Purpose"], [["Orders.csv", "500", "10", "Order-level transactions, amounts, payments, delivery time, and ratings"], ["Restaurants.csv", "20", "5", "Restaurant master data, cuisine, zone, and category"]], [1.45, .6, .7, 3.75])
    heading(doc, "5.1 Orders Attributes", 2)
    add_table(doc, ["Column", "Description"], [["Order ID", "Unique order identifier"], ["Customer Name", "Customer name"], ["Restaurant ID", "Foreign key to restaurant master"], ["Order Date", "Date and time of order"], ["Quantity of Items", "Number of items ordered"], ["Order Amount", "Total bill value"], ["Payment Mode", "Mode used for payment"], ["Delivery Time Taken (mins)", "Delivery duration in minutes"], ["Customer Rating-Food", "Food rating, range 1–5"], ["Customer Rating-Delivery", "Delivery rating, range 1–5"]], [2.5, 4.0])
    heading(doc, "5.2 Restaurants Attributes", 2)
    add_table(doc, ["Column", "Description"], [["RestaurantID", "Unique restaurant identifier"], ["RestaurantName", "Restaurant name"], ["Cuisine", "Cuisine type"], ["Zone", "Operational zone"], ["Category", "Restaurant category: Pro or Ordinary"]], [2.5, 4.0])
    page_break(doc)

    heading(doc, "Chapter 6  Methodology")
    heading(doc, "6.1 Extract", 2)
    add_text(doc, "The pipeline locates source files from data/raw or the supplied Downloads directory. Pandas reads the CSV files into tabular data frames.")
    heading(doc, "6.2 Transform and Validate", 2)
    add_bullets(doc, ["Convert headers to snake_case for consistent programmatic access.", "Strip extra spaces from textual fields.", "Convert order_date to a datetime value and numerical fields to numeric types.", "Standardize payment-mode labels such as UPI, Credit Card, Debit Card, Cash, and Cash on Delivery.", "Remove repeated order IDs and duplicate records.", "Reject records with invalid or missing key values, non-positive quantity/amount/delivery time, or ratings outside 1–5.", "Standardize restaurant names, cuisine, zone, and category values."])
    heading(doc, "6.3 Join and Feature Engineering", 2)
    add_text(doc, "Validated orders are left-joined to restaurant master data on restaurant_id. The pipeline derives order_month, day_of_week, is_weekend_order, delivery_speed, and average_rating. Delivery speed is defined as Fast (up to 20 minutes), Normal (21–40 minutes), and Slow (more than 40 minutes).")
    heading(doc, "6.4 Load", 2)
    add_text(doc, "The final curated dataset is written to processed_orders.csv. The same data is loaded into SQLite as curated_orders, and restaurant master data is stored as restaurants. A JSON validation report documents processing counts and quality outcomes.")
    page_break(doc)

    heading(doc, "Chapter 7  Implementation")
    add_text(doc, "The implementation is organized into two Python modules. etl_pipeline.py contains reusable cleaning, feature engineering, output loading, and command-line logic. dashboard.py loads curated data and presents it using Streamlit and Plotly. The environment dependencies are Pandas, Plotly, and Streamlit.")
    heading(doc, "7.1 Key Implementation Functions", 2)
    add_table(doc, ["Function", "Responsibility"], [["resolve_source", "Finds raw CSV files from an explicit path, project data folder, or Downloads folder"], ["clean_orders", "Renames, parses, standardizes, removes duplicates, and validates order records"], ["clean_restaurants", "Normalizes restaurant master data and removes duplicate IDs"], ["add_features", "Creates month, day, weekend, speed, and average-rating fields"], ["run_pipeline", "Coordinates extract, transform, join, feature engineering, CSV, SQLite, and report load"], ["get_data", "Reads the curated dataset for the dashboard and runs ETL if required"]], [1.6, 4.9])
    heading(doc, "7.2 Reproducibility", 2)
    add_text(doc, "The project includes requirements.txt and README.md. A virtual environment can be created and dependencies installed before executing the pipeline or dashboard. This avoids dependence on manual edits and allows the same source inputs to produce the same curated outputs.")
    page_break(doc)

    heading(doc, "Chapter 8  Dashboard")
    add_text(doc, "The Streamlit dashboard provides an accessible decision-support interface over curated data. The sidebar supports filtering by zone, cuisine, payment mode, and order-date range. All KPIs and charts update according to the selected filters.")
    add_table(doc, ["Dashboard element", "Purpose"], [["KPI cards", "Show order count, total revenue, average delivery time, and average rating"], ["Top restaurants chart", "Ranks restaurants by order count and shows associated revenue"], ["Payment-mode chart", "Displays payment-method share"], ["Order amount histogram", "Shows distribution of order value"], ["Delivery-time box chart", "Compares delivery time across speed tiers"], ["Cuisine and zone charts", "Compare cuisine demand and zone revenue"], ["Curated data table", "Provides row-level inspection and filtered CSV download"]], [2.15, 4.35])
    heading(doc, "8.1 Dashboard Value", 2)
    add_text(doc, "The dashboard enables restaurant managers and analysts to move from raw records to a filtered, visual interpretation without writing SQL or data-processing code. It also ensures that displayed values are based on validated and enriched data rather than raw source files.")
    page_break(doc)

    heading(doc, "Chapter 9  Testing")
    add_text(doc, "Testing combined pipeline execution, output checks, compilation, and in-process dashboard validation. The dashboard was verified with Streamlit's testing interface; it rendered without exceptions and produced four KPI metrics and six Plotly charts.")
    add_table(doc, ["Test case", "Expected result", "Outcome"], [["Run ETL on supplied CSVs", "Curated output and report generated", "Pass"], ["Validate data quality rules", "Invalid rows, duplicates, and missing values reported", "Pass"], ["Check join completeness", "All order restaurant IDs match master data", "Pass"], ["Inspect processed CSV", "500 rows and engineered features present", "Pass"], ["Compile Python modules", "No syntax errors", "Pass"], ["Dashboard render test", "No exceptions; 4 metrics and 6 charts", "Pass"]], [2.1, 3.2, 1.2])
    heading(doc, "9.1 Validation Evidence", 2)
    add_text(doc, "The persisted validation report records source_orders = 500, source_restaurants = 20, curated_orders = 500, rejected_orders = 0, unmatched_restaurants = 0, duplicate_order_ids = 0, and missing_values_after_cleaning = 0.")
    page_break(doc)

    heading(doc, "Chapter 10  Results")
    add_text(doc, "The ETL run successfully transformed the source files into a unified analytical dataset. All 500 source orders were retained because each record passed the defined validation rules and all restaurant identifiers matched the master table. The curated dataset contains 19 columns: the original cleaned order fields, restaurant attributes, and five engineered fields.")
    add_table(doc, ["Measure", "Result"], [["Source orders", "500"], ["Source restaurants", "20"], ["Curated orders", "500"], ["Rejected orders", "0"], ["Unmatched restaurants", "0"], ["Duplicate order IDs", "0"], ["Missing values after cleaning", "0"], ["Curated dataset columns", "19"]], [3.6, 2.9])
    heading(doc, "10.1 Interpretation", 2)
    add_text(doc, "The results indicate that the supplied dataset is structurally clean after standardization. Although no rows were rejected in this run, the validation logic is still important because it protects the reporting layer if future files contain bad dates, invalid ratings, non-positive amounts, duplicate identifiers, or unrecognized restaurant relationships.")
    page_break(doc)

    heading(doc, "Chapter 11  Conclusion")
    add_text(doc, "The Restaurant Orders Data Curation project meets its objective of creating a complete pipeline from raw data to dashboard. It demonstrates practical data engineering steps: source discovery, profiling-oriented cleaning, business-rule validation, relational integration, feature engineering, multi-format loading, and interactive analytics. The generated CSV, SQLite database, and validation report provide reliable artifacts for future reporting or analysis.")
    add_text(doc, "By separating raw sources from curated outputs, the solution establishes a dependable foundation for restaurant analytics. The reusable design makes it simple to rerun the pipeline when data is updated.")
    page_break(doc)

    heading(doc, "Chapter 12  Future Scope")
    add_bullets(doc, ["Add automated data-profiling summaries and anomaly detection.", "Schedule the ETL process with a workflow orchestrator such as Airflow or Prefect.", "Move from SQLite to PostgreSQL or a cloud warehouse for larger datasets.", "Add incremental loads, audit columns, and historical snapshots.", "Integrate live restaurant, customer, and delivery-partner data through APIs.", "Build predictive models for delivery time, customer rating, and revenue forecasting.", "Deploy the dashboard on Streamlit Community Cloud, Azure, AWS, or another hosting platform.", "Add role-based access, automated tests, and CI/CD pipelines."])
    page_break(doc)

    heading(doc, "References")
    references = [
        "Kimball, R. and Caserta, J. (2013). The Data Warehouse ETL Toolkit: Practical Techniques for Extracting, Cleaning, Conforming, and Delivering Data. 2nd ed. Wiley.",
        "McKinney, W. (2022). Python for Data Analysis. 3rd ed. O’Reilly Media.",
        "Pandas Development Team. pandas Documentation. Available at: https://pandas.pydata.org/docs/ (Accessed: 14 July 2026).",
        "SQLite Consortium. SQLite Documentation. Available at: https://www.sqlite.org/docs.html (Accessed: 14 July 2026).",
        "Streamlit Inc. Streamlit Documentation. Available at: https://docs.streamlit.io/ (Accessed: 14 July 2026).",
        "Plotly Technologies Inc. Plotly Python Documentation. Available at: https://plotly.com/python/ (Accessed: 14 July 2026).",
    ]
    for ref in references:
        add_text(doc, ref, after=7)
    page_break(doc)

    heading(doc, "Appendices")
    heading(doc, "Appendix A  Project Folder Structure", 2)
    add_text(doc, "restaurant-orders-etl/\n├── etl_pipeline.py\n├── dashboard.py\n├── requirements.txt\n├── README.md\n├── data/processed/\n│   ├── processed_orders.csv\n│   ├── restaurant_orders.db\n│   └── validation_report.json\n└── reports/", size=9)
    page_break(doc)
    add_code(doc, ROOT / "etl_pipeline.py", "Appendix B  Complete ETL Pipeline Source Code")
    page_break(doc)
    add_code(doc, ROOT / "dashboard.py", "Appendix C  Complete Streamlit Dashboard Source Code")
    page_break(doc)
    heading(doc, "Appendix D  Requirements and Execution Commands", 2)
    add_text(doc, "requirements.txt", bold=True)
    for line in (ROOT / "requirements.txt").read_text().splitlines():
        add_text(doc, line, size=9, after=1)
    add_text(doc, "Execution commands", bold=True, after=4)
    for command in ["cd restaurant-orders-etl", "python3 -m venv .venv", "source .venv/bin/activate", "pip install -r requirements.txt", "python etl_pipeline.py --orders ~/Downloads/Orders.csv --restaurants ~/Downloads/Restaurants.csv", "streamlit run dashboard.py"]:
        p = doc.add_paragraph()
        r = p.add_run(command)
        set_font(r, size=9, name="Courier New")
        p.paragraph_format.space_after = Pt(2)

    doc.core_properties.title = "Restaurant Orders Data Curation Project Report"
    doc.core_properties.author = "[STUDENT NAME]"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
